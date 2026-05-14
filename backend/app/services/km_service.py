"""KM 服務：文字擷取、Chunking、Embedding、向量檢索

設計：
  - 使用 pgvector 做向量搜尋（cosine similarity）
  - Embedding model：由 tenant_configs 鎖定（預設 Gemini text-embedding-004，768 維）
    換模型需走遷移流程（清空向量索引、re-embed、version +1）
  - 支援 PDF（pypdf）、純文字/Markdown（UTF-8）
  - chunk_size=1500 字元，overlap=200 字元
"""
import io
import logging
import time
from datetime import datetime, timezone

import litellm
import sqlalchemy as sa
from sqlalchemy.orm import Session

from app.models.km_chunk import KmChunk
from app.models.km_document import KmDocument
from app.models.km_query_log import KmQueryLog
from app.models.km_knowledge_base import KmKnowledgeBase
from app.models.llm_provider_config import LLMProviderConfig
from app.models.tenant_config import TenantConfig
from app.core.encryption import decrypt_api_key
from app.services.agent_usage import log_agent_usage

logger = logging.getLogger(__name__)

EMBEDDING_DIM = 768
CHUNK_SIZE = 1500
CHUNK_OVERLAP = 200

# Embedding model litellm 前綴對應（provider → litellm model 字串前綴）
# 各文件類型的 chunking 策略與檢索 top_k
# chunk_size 越小 → top_k 越大（總 context 字元數約 3000–6000）
CHUNK_STRATEGIES: dict[str, dict] = {
    "faq":       {"chunk_size": 300,  "overlap": 50,  "top_k": 12},
    "spec":      {"chunk_size": 400,  "overlap": 80,  "top_k": 12},
    "chat":      {"chunk_size": 500,  "overlap": 50,  "top_k": 10},
    "policy":    {"chunk_size": 800,  "overlap": 150, "top_k": 6},
    "article":   {"chunk_size": 1500, "overlap": 200, "top_k": 4},
    "reference": {"chunk_size": 60000, "overlap": 0,   "top_k": 10},
}
DOC_TYPES = frozenset(CHUNK_STRATEGIES.keys())


def _strategy(doc_type: str) -> dict:
    """取得文件類型對應策略，未知類型 fallback 到 article。"""
    return CHUNK_STRATEGIES.get(doc_type, CHUNK_STRATEGIES["article"])


# ──────────────────────────────────────────────────────────────────────────────
# 文字擷取
# ──────────────────────────────────────────────────────────────────────────────


def extract_text(file_bytes: bytes, content_type: str | None, filename: str) -> str:
    """從 PDF、Word 或文字檔擷取純文字。"""
    ct = (content_type or "").lower()
    name = (filename or "").lower()

    if "pdf" in ct or name.endswith(".pdf"):
        return _extract_pdf(file_bytes)

    if (
        "wordprocessingml" in ct
        or "msword" in ct
        or name.endswith(".docx")
        or name.endswith(".doc")
    ):
        return _extract_docx(file_bytes)

    # 圖片：暫不支援 OCR，回傳空字串讓上層處理
    if ct.startswith("image/") or any(name.endswith(ext) for ext in (".jpg", ".jpeg", ".png", ".gif", ".webp")):
        return ""

    # 純文字 / Markdown
    try:
        return file_bytes.decode("utf-8").strip()
    except UnicodeDecodeError:
        return file_bytes.decode("latin-1", errors="replace").strip()


def _extract_docx(file_bytes: bytes) -> str:
    """從 Word .docx 擷取純文字（段落 + 表格）。"""
    import io
    try:
        import docx as python_docx
    except ImportError:
        logger.warning("python-docx 未安裝，無法擷取 Word 文件內容")
        return ""
    try:
        doc = python_docx.Document(io.BytesIO(file_bytes))
        parts: list[str] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)
        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts)
    except Exception as e:
        logger.warning("Word 文件擷取失敗：%s", e)
        return ""


def _extract_pdf(file_bytes: bytes) -> str:
    """從 PDF 擷取純文字。
    策略：pdfplumber（主）→ pypdf（備援）。
    pdfplumber 對表格式、多欄 PDF（如 Apple 規格書）提取能力較強。
    提取後會清洗掉 URL、頁碼、時間戳記等雜訊行。
    """
    text = _extract_pdf_pdfplumber(file_bytes)
    if len(text.strip()) < 100:
        text_pypdf = _extract_pdf_pypdf(file_bytes)
        if len(text_pypdf) > len(text):
            text = text_pypdf
    return _clean_pdf_text(text)


def _clean_pdf_text(text: str) -> str:
    """清洗 PDF 提取文字：移除 URL、頁碼行、時間戳記等雜訊。"""
    import re

    clean_lines: list[str] = []
    for line in text.splitlines():
        s = line.strip()
        if not s:
            clean_lines.append("")
            continue
        # 過濾純 URL 行
        if re.match(r"^https?://\S+$", s):
            continue
        # 過濾 URL + 頁碼行（如：https://support.apple.com/... 4/11）
        if re.match(r"^https?://\S+\s+\d+/\d+$", s):
            continue
        # 過濾「日期 時間 標題 URL」混合行（如：2026/4/17 下午3:23 iPhone 17 Pro - 技術規格...）
        if re.match(r"^\d{4}/\d{1,2}/\d{1,2}\s", s):
            continue
        # 過濾頁碼行（如：1/11、2/11）
        if re.match(r"^\d+/\d+$", s):
            continue
        clean_lines.append(line)

    # 壓縮連續空白行為單一空行
    result = re.sub(r"\n{3,}", "\n\n", "\n".join(clean_lines))
    return result.strip()


def _table_to_markdown(table: list[list]) -> str:
    """將 pdfplumber 二維陣列轉成 Markdown 表格字串。"""
    rows = [[str(cell or "").strip() for cell in row] for row in table if any(cell for cell in row)]
    if not rows:
        return ""
    header = rows[0]
    body = rows[1:]
    sep = ["---"] * len(header)
    lines = [
        "| " + " | ".join(header) + " |",
        "| " + " | ".join(sep) + " |",
    ] + ["| " + " | ".join(row) + " |" for row in body]
    return "\n".join(lines)


def _extract_pdf_pdfplumber(file_bytes: bytes) -> str:
    try:
        import pdfplumber  # type: ignore[import-untyped]

        parts: list[str] = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_parts: list[str] = []

                # 1. 找出所有表格及其 bbox，轉成 Markdown
                table_objs = page.find_tables()
                table_bboxes = [t.bbox for t in table_objs]
                for t in table_objs:
                    md = _table_to_markdown(t.extract())
                    if md:
                        page_parts.append(md)

                # 2. 萃取表格區域以外的純文字，避免重複
                if table_bboxes:
                    words = page.extract_words(x_tolerance=2, y_tolerance=3)
                    non_table = []
                    for w in words:
                        wx0, wy0, wx1, wy1 = w["x0"], w["top"], w["x1"], w["bottom"]
                        in_table = any(
                            wx0 >= bx0 and wy0 >= by0 and wx1 <= bx1 and wy1 <= by1
                            for bx0, by0, bx1, by1 in table_bboxes
                        )
                        if not in_table:
                            non_table.append(w["text"])
                    plain = " ".join(non_table).strip()
                else:
                    plain = (page.extract_text(x_tolerance=2, y_tolerance=3) or "").strip()

                if plain:
                    page_parts.insert(0, plain)  # 純文字置於表格前

                if page_parts:
                    parts.append("\n\n".join(page_parts))

        return "\n\n".join(parts)
    except Exception as e:
        logger.warning("pdfplumber 擷取失敗: %s", e)
        return ""


def _collapse_spaced_text(text: str) -> str:
    """修正 fpdf2 CIDFont 產生的「每字元間有空格」問題。
    偵測到大量字元間空格時，移除所有非換行的字間空格。
    """
    import re
    if not text:
        return text
    # 計算非換行字元數和空格數，若空格率 > 60% 視為 spaced-out text
    non_newline = text.replace("\n", "")
    char_count = len(non_newline.replace(" ", ""))
    space_count = non_newline.count(" ")
    if char_count > 0 and space_count / max(char_count, 1) > 0.6:
        # 移除相鄰可見字元之間的單一空格
        text = re.sub(r"(?<=\S) (?=\S)", "", text)
        # 壓縮殘留的多餘空格
        text = re.sub(r"  +", " ", text)
    return text


def _extract_pdf_pypdf(file_bytes: bytes) -> str:
    try:
        from pypdf import PdfReader  # type: ignore[import-untyped]

        reader = PdfReader(io.BytesIO(file_bytes))
        parts: list[str] = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                parts.append(_collapse_spaced_text(text.strip()))
        return "\n\n".join(parts)
    except Exception as e:
        logger.warning("pypdf 擷取失敗: %s", e)
        return ""


# ──────────────────────────────────────────────────────────────────────────────
# Chunking
# ──────────────────────────────────────────────────────────────────────────────

import re as _re


def _chunk_sliding(text: str, chunk_size: int, overlap: int) -> list[str]:
    """滑動視窗切分（基礎實作）。"""
    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        piece = text[start:end].strip()
        if piece:
            chunks.append(piece)
        if end >= len(text):
            break
        start = end - overlap
    return chunks


def detect_faq_format(text: str) -> bool:
    """判斷文字是否包含可辨識的 FAQ 結構（Layer 1 或 Layer 2）。
    僅 Q 前綴或 ●/○ bullet 格式視為合格 FAQ；
    純段落或無結構文字不算（那是任何文章都能通過的 fallback）。

    至少需有 2 組 Q/A 對才算有效。
    """
    # Layer 1：Q 語意前綴（Q: / 問: / 問題: 等）
    q_prefix = _re.compile(
        r'(?m)^[ \t]*(?:Q[：:]|Question\s*[：:]|問[：:]|問題[：:])',
    )
    if len(q_prefix.findall(text)) >= 2:
        return True

    # Layer 2：●/○ bullet 格式
    bullet_pattern = _re.compile(r'(?:^|\n)(?=\s*[●•]\s)', _re.MULTILINE)
    bullet_parts = [p.strip() for p in bullet_pattern.split(text) if p.strip()]
    if len(bullet_parts) >= 2:
        return True

    return False


def _chunk_faq(text: str) -> list[str]:
    """FAQ 專用切割：每個 Q&A 對成一個完整 chunk。

    切割策略（依序嘗試）：
      1. 以明確語意前綴（Q: / 問: / 問題: 等）為邊界，收集到下一個 Q 前綴為止
         → 保證答案中的編號步驟不被誤切
      2. ●/○ bullet 格式（● 問題 + ○ 答案）
      3. 雙換行段落切
      4. fallback 滑動視窗（chunk_size=300）
    """
    # Layer 1：以 Q 語意前綴為邊界（只切 Q 不切 A，避免步驟被拆散）
    # 支援格式：Q: / Q： / Question: / 問: / 問： / 問題: / 問題：
    q_prefix = _re.compile(
        r'(?m)^[ \t]*(?:Q[：:]|Question\s*[：:]|問[：:]|問題[：:])',
    )
    q_positions = [m.start() for m in q_prefix.finditer(text)]
    if len(q_positions) >= 2:
        chunks = []
        for i, start in enumerate(q_positions):
            end = q_positions[i + 1] if i + 1 < len(q_positions) else len(text)
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
        logger.debug("FAQ chunking：Q-prefix，%d 對", len(chunks))
        return chunks

    # Layer 2：● 問題 / ○ 答案 的 bullet 格式
    bullet_pattern = _re.compile(r'(?:^|\n)(?=\s*[●•]\s)', _re.MULTILINE)
    bullet_parts = [p.strip() for p in bullet_pattern.split(text) if p.strip()]
    if len(bullet_parts) >= 2:
        logger.debug("FAQ chunking：● bullet 格式，%d 對", len(bullet_parts))
        return bullet_parts

    # Layer 3：按段落切（雙換行）
    paragraphs = [p.strip() for p in _re.split(r'\n{2,}', text) if p.strip()]
    if len(paragraphs) >= 2:
        logger.debug("FAQ chunking：段落模式，%d 段", len(paragraphs))
        return paragraphs

    # Layer 4：滑動視窗 fallback
    logger.debug("FAQ chunking：fallback 滑動視窗")
    return _chunk_sliding(text, chunk_size=300, overlap=50)


def chunk_text(
    text: str,
    doc_type: str = "article",
    chunk_size: int | None = None,
    overlap: int | None = None,
) -> list[str]:
    """文件切分入口：依 doc_type 選擇切分策略。
    - faq：Q&A 感知切割（每對獨立 chunk）
    - 其他：依 CHUNK_STRATEGIES 的 chunk_size/overlap 滑動視窗
    """
    text = text.strip()
    if not text:
        return []

    if doc_type == "faq":
        return _chunk_faq(text)

    s = _strategy(doc_type)
    size = chunk_size if chunk_size is not None else s["chunk_size"]
    ov = overlap if overlap is not None else s["overlap"]
    return _chunk_sliding(text, chunk_size=size, overlap=ov)


# ──────────────────────────────────────────────────────────────────────────────
# Embedding
# ──────────────────────────────────────────────────────────────────────────────


def _get_embed_params(db: Session, tenant_id: str) -> tuple[str, str, str | None, str | None] | None:
    """
    從 tenant_configs 讀取鎖定的 embedding 設定，
    再從 llm_provider_configs 取對應 provider 的 API key。
    回傳 (provider, model, api_key, api_base) 或 None（設定不存在）。
    model 直接使用 tenant_configs.embedding_model，不自動加前綴。
    """
    tc = db.query(TenantConfig).filter(TenantConfig.tenant_id == tenant_id).first()
    if not tc:
        return None

    provider = tc.embedding_provider
    model_name = tc.embedding_model  # 直接使用，不加前綴

    # 尚未設定 embedding，回傳 None
    if not provider or not model_name:
        return None

    # 取對應 provider 的 API key
    provider_cfg = (
        db.query(LLMProviderConfig)
        .filter(
            LLMProviderConfig.tenant_id == tenant_id,
            LLMProviderConfig.provider == provider,
            LLMProviderConfig.is_active.is_(True),
        )
        .order_by(LLMProviderConfig.id)
        .first()
    )

    api_key: str | None = None
    api_base: str | None = None

    if provider_cfg:
        api_base = provider_cfg.api_base_url or None
        if provider_cfg.api_key_encrypted:
            try:
                api_key = decrypt_api_key(provider_cfg.api_key_encrypted)
            except ValueError:
                api_key = None

    # local provider 特殊處理：api_key 可為任意字串；api_base 預設 localhost:11434
    if provider == "local":
        api_key = api_key or "local"
        if not api_base:
            api_base = "http://localhost:11434"
    elif not api_key:
        return None

    return provider, model_name, api_key, api_base


def _lock_embedding_config(db: Session, tenant_id: str) -> None:
    """第一次寫入 embedding 向量時鎖定 tenant 的 embedding 設定（冪等）。"""
    db.query(TenantConfig).filter(
        TenantConfig.tenant_id == tenant_id,
        TenantConfig.embedding_locked_at.is_(None),
    ).update(
        {"embedding_locked_at": datetime.now(timezone.utc)},
        synchronize_session=False,
    )
    db.commit()




def embed_texts_sync(
    texts: list[str],
    model: str,
    api_key: str,
    provider: str = "openai",
    api_base: str | None = None,
    task_type: str = "retrieval_document",
) -> tuple[list[list[float]], int | None]:
    """同步呼叫 embedding API，回傳 (向量清單, prompt_tokens)。
    所有 provider 統一走 LiteLLM：
      - gemini → LiteLLM gemini/ 前綴（task_type 透過 extra_body 傳遞）
      - openai → LiteLLM
      - local  → LiteLLM + ollama api_base
    model 直接使用設定值，不自動加前綴。
    task_type：retrieval_document（上傳文件）或 retrieval_query（搜尋查詢）
    """
    if provider == "gemini":
        litellm_model = model if model.startswith("gemini/") else f"gemini/{model}"
        kwargs: dict = dict(
            model=litellm_model,
            input=texts,
            api_key=api_key,
            timeout=15,
            extra_body={"task_type": task_type, "output_dimensionality": EMBEDDING_DIM},
        )
        response = litellm.embedding(**kwargs)
        prompt_tokens: int | None = getattr(getattr(response, "usage", None), "prompt_tokens", None)
        return [item["embedding"] for item in response.data], prompt_tokens

    if provider == "local":
        # Ollama embedding：LiteLLM 需要 ollama/ 前綴；api_base 不加 /v1
        litellm_model = model if model.startswith("ollama/") else f"ollama/{model}"
        kwargs = dict(model=litellm_model, input=texts, api_key=api_key or "local", timeout=15)
        if api_base:
            kwargs["api_base"] = api_base.rstrip("/")
        response = litellm.embedding(**kwargs)
        prompt_tokens = getattr(getattr(response, "usage", None), "prompt_tokens", None)
        return [item["embedding"] for item in response.data], prompt_tokens

    kwargs = dict(model=model, input=texts, api_key=api_key, timeout=15)
    if api_base:
        kwargs["api_base"] = api_base
    if provider == "openai" and "text-embedding-3-small" in model:
        kwargs["dimensions"] = EMBEDDING_DIM
    response = litellm.embedding(**kwargs)
    prompt_tokens = getattr(getattr(response, "usage", None), "prompt_tokens", None)
    return [item["embedding"] for item in response.data], prompt_tokens


# ──────────────────────────────────────────────────────────────────────────────
# 文件處理管線
# ──────────────────────────────────────────────────────────────────────────────


def process_document(
    doc_id: int,
    file_bytes: bytes,
    content_type: str | None,
    filename: str,
    db: Session,
    tenant_id: str,
    doc_type: str = "article",
    agent_id: str = "knowledge",
    user_id: int | None = None,
) -> None:
    """完整管線：文字擷取 → 切分 → Embedding → 寫入 km_chunks。
    同步執行（於上傳請求中直接處理，適合初期小量文件場景）。
    doc_type 決定 chunk_size / overlap（見 CHUNK_STRATEGIES）。
    """
    doc = db.get(KmDocument, doc_id)
    if not doc:
        return

    try:
        doc.status = "processing"
        db.commit()

        # 1. 擷取文字
        text = extract_text(file_bytes, content_type, filename)
        if not text.strip():
            doc.status = "error"
            doc.error_message = "無法從檔案擷取文字內容（PDF 可能加密，或為純圖片 PDF）"
            db.commit()
            return

        # 2. 依文件類型切分 chunks
        chunks = chunk_text(text, doc_type=doc_type)
        if not chunks:
            doc.status = "error"
            doc.error_message = "文字切分失敗"
            db.commit()
            return

        # 3. 讀取 tenant 鎖定的 Embedding 設定
        embed_params = _get_embed_params(db, tenant_id)
        if not embed_params:
            doc.status = "error"
            doc.error_message = "未設定 Embedding Provider 或 API Key，無法產生 Embedding。請在管理介面設定 Gemini、OpenAI 或 Local provider。"
            db.commit()
            return
        embed_provider, embed_model, embed_key, embed_base = embed_params
        logger.info("KM embed 使用 provider=%s model=%s (doc_id=%d)", embed_provider, embed_model, doc_id)

        # 4. 批次 Embedding（每批 100 筆）
        # FAQ/其他類型一律對完整 chunk 內容（Q+A 全文）做 embedding，提高語意召回率。
        # BM25（content_tsv）才是 FAQ Q-only；兩者設計不同，RRF 時互補。
        embed_texts = chunks
        all_embeddings: list[list[float]] = []
        batch_size = 100
        embed_started = time.monotonic()
        embed_status = "success"
        total_prompt_tokens: int | None = None
        try:
            for i in range(0, len(embed_texts), batch_size):
                batch = embed_texts[i : i + batch_size]
                batch_vectors, batch_tokens = embed_texts_sync(batch, model=embed_model, api_key=embed_key, provider=embed_provider, api_base=embed_base)
                all_embeddings.extend(batch_vectors)
                if batch_tokens is not None:
                    total_prompt_tokens = (total_prompt_tokens or 0) + batch_tokens
        except Exception:
            embed_status = "error"
            raise
        finally:
            log_agent_usage(
                db=db,
                agent_type=agent_id,
                tenant_id=tenant_id,
                user_id=user_id,
                model=embed_model,
                prompt_tokens=total_prompt_tokens,
                total_tokens=total_prompt_tokens,
                latency_ms=int((time.monotonic() - embed_started) * 1000),
                status=embed_status,
            )

        # 5. 寫入 km_chunks，並鎖定 embedding config
        # BM25 tsvector 只對短、語意集中的類型有意義（faq/spec），其餘不建 index
        # FAQ type：content_tsv 索引 Q+A 全文（與 embedding 一致），
        # 確保 tag、答案術語等關鍵字都能被 BM25 搜到；
        # 語意精準度由 embedding（ALPHA=0.7）主導，BM25 負責關鍵字覆蓋。
        BM25_DOC_TYPES = {"faq", "spec"}
        build_tsv = doc_type in BM25_DOC_TYPES
        for idx, (chunk_content, embedding) in enumerate(zip(chunks, all_embeddings)):
            tsv_text = chunk_content
            km_chunk = KmChunk(
                document_id=doc_id,
                chunk_index=idx,
                content=chunk_content,
                embedding=embedding,
                content_tsv=(
                    db.execute(
                        sa.text("SELECT to_tsvector('public.cjk', :t)"),
                        {"t": tsv_text},
                    ).scalar()
                    if build_tsv else None
                ),
                metadata_={"filename": filename, "chunk_index": idx},
            )
            db.add(km_chunk)

        _lock_embedding_config(db, tenant_id)

        doc.status = "ready"
        doc.chunk_count = len(chunks)
        db.commit()

        logger.info(
            "KM 文件 id=%d 處理完成：%d chunks，檔名=%s",
            doc_id,
            len(chunks),
            filename,
        )

    except Exception as e:
        logger.exception("KM 文件 id=%d 處理失敗", doc_id)
        try:
            doc.status = "error"
            doc.error_message = str(e)[:1000]
            db.commit()
        except Exception:
            pass


# ──────────────────────────────────────────────────────────────────────────────
# 向量檢索
# ──────────────────────────────────────────────────────────────────────────────


def km_retrieve_sync(
    query: str,
    db: Session,
    tenant_id: str,
    user_id: int = 0,
    top_k: int | None = None,
    selected_doc_ids: list[int] | None = None,
    knowledge_base_id: int | None = None,
    knowledge_base_ids: list[int] | None = None,
    skip_scope_check: bool = False,
    agent_id: str = "knowledge",
) -> list[KmChunk]:
    """同步向量檢索：query → embedding → cosine similarity 找 top-K chunks。

    top_k：若未指定，依 KB 內文件的 doc_type 自動決定（取最大 top_k 以保守覆蓋）。
    存取範圍（新邏輯，依 KB scope 決定）：
      - KB scope='company' → 同 tenant 全員可搜尋
      - KB scope='personal' → 只有 KB 建立者（created_by）可搜尋
      - 無 KB（knowledge_base_id = NULL）→ 沿用舊文件 owner 邏輯
    若提供 knowledge_base_id，只在該知識庫的文件中搜尋。
    若提供 knowledge_base_ids（非空 list），在多個知識庫中聯合搜尋（Bot 多 KB 模式）。
    若提供 selected_doc_ids（非空），則只在指定文件中搜尋。
    skip_scope_check=True：跳過 scope 過濾（Widget 公開存取用）。
    """
    # knowledge_base_ids 多 KB 模式：knowledge_base_id 退為 None
    if knowledge_base_ids:
        knowledge_base_id = None
    embed_params = _get_embed_params(db, tenant_id)
    if not embed_params:
        logger.warning("KM 檢索失敗：tenant_id=%s 無可用 Embedding provider", tenant_id)
        return []
    embed_provider, embed_model, embed_key, embed_base = embed_params

    # 動態決定 top_k：依 KB 內文件類型取最大值（混合類型保守取多）
    if top_k is None:
        try:
            q_types = db.query(KmDocument.doc_type).filter(
                KmDocument.tenant_id == tenant_id,
                KmDocument.status == "ready",
            )
            if knowledge_base_ids:
                q_types = q_types.filter(KmDocument.knowledge_base_id.in_(knowledge_base_ids))
            elif knowledge_base_id is not None:
                q_types = q_types.filter(KmDocument.knowledge_base_id == knowledge_base_id)
            elif selected_doc_ids:
                q_types = q_types.filter(KmDocument.id.in_(selected_doc_ids))
            doc_types = [row[0] for row in q_types.distinct().all()]
            top_k = max(
                (_strategy(dt)["top_k"] for dt in doc_types),
                default=8,
            )
        except Exception:
            top_k = 8
        logger.debug("km_retrieve top_k=%d (doc_types=%s)", top_k, doc_types if doc_types else "unknown")

    embed_started = time.monotonic()
    embed_status = "success"
    query_prompt_tokens: int | None = None
    try:
        vectors, query_prompt_tokens = embed_texts_sync(
            [query],
            model=embed_model,
            api_key=embed_key,
            provider=embed_provider,
            api_base=embed_base,
            task_type="retrieval_query",
        )
        query_embedding = vectors[0]
    except Exception as e:
        embed_status = "error"
        logger.warning("KM embedding 失敗: %s", e)
        return []
    finally:
        log_agent_usage(
            db=db,
            agent_type=agent_id,
            tenant_id=tenant_id,
            user_id=user_id if user_id else None,
            model=embed_model,
            prompt_tokens=query_prompt_tokens,
            total_tokens=query_prompt_tokens,
            latency_ms=int((time.monotonic() - embed_started) * 1000),
            status=embed_status,
        )

    try:
        # ── 共用 filter 條件 ──────────────────────────────────
        def _base_filters():
            filters = [
                KmDocument.tenant_id == tenant_id,
                KmDocument.status == "ready",
            ]
            if not skip_scope_check:
                # 新邏輯：依 KB scope 判斷可見性
                #   company KB → 全 tenant 員工可見
                #   personal KB → 只有 KB 建立者（created_by）可搜尋
                #   無 KB（knowledge_base_id = NULL）→ 沿用舊文件 owner 邏輯
                from sqlalchemy import or_, and_
                filters.append(
                    or_(
                        # 無 KB 的舊文件：沿用 owner 邏輯
                        and_(
                            KmDocument.knowledge_base_id.is_(None),
                            KmDocument.owner_user_id == user_id,
                        ),
                        # 有 KB 的文件：依 KB scope 判斷
                        and_(
                            KmDocument.knowledge_base_id.isnot(None),
                            or_(
                                KmKnowledgeBase.scope == "company",
                                KmKnowledgeBase.created_by == user_id,
                            ),
                        ),
                    )
                )
            if knowledge_base_ids:
                filters.append(KmDocument.knowledge_base_id.in_(knowledge_base_ids))
            elif knowledge_base_id is not None:
                filters.append(KmDocument.knowledge_base_id == knowledge_base_id)
            if selected_doc_ids:
                filters.append(KmDocument.id.in_(selected_doc_ids))
            return filters

        fetch_k = top_k * 3  # 各自多撈一些，合併後再取 top_k

        # ── 向量搜尋（cosine distance 升序 = 相似度降序）──────
        vector_rows = (
            db.query(KmChunk)
            .join(KmDocument, KmChunk.document_id == KmDocument.id)
            .outerjoin(KmKnowledgeBase, KmDocument.knowledge_base_id == KmKnowledgeBase.id)
            .filter(*_base_filters())
            .order_by(KmChunk.embedding.cosine_distance(query_embedding))
            .limit(fetch_k)
            .all()
        )

        # ── BM25 全文搜尋（OR-tsquery：2-gram lexeme 用 OR 連接）──
        bm25_rows: list[KmChunk] = []
        try:
            tsv_raw = db.execute(
                sa.text("SELECT to_tsvector('public.cjk', :q)"), {"q": query}
            ).scalar()
            or_tsquery: str | None = None
            if tsv_raw:
                lexemes = [
                    f"'{part.split(':')[0].strip(chr(39))}'"
                    for part in str(tsv_raw).split()
                    if part
                ]
                if lexemes:
                    or_tsquery = " | ".join(lexemes)

            if or_tsquery:
                tsq = sa.func.to_tsquery("public.cjk", or_tsquery)
                bm25_q = (
                    db.query(KmChunk)
                    .join(KmDocument, KmChunk.document_id == KmDocument.id)
                    .outerjoin(KmKnowledgeBase, KmDocument.knowledge_base_id == KmKnowledgeBase.id)
                    .filter(
                        *_base_filters(),
                        KmChunk.content_tsv.op("@@")(tsq),
                    )
                    .order_by(sa.func.ts_rank(KmChunk.content_tsv, tsq, 1).desc())
                    .limit(fetch_k)
                )
                bm25_rows = bm25_q.all()
        except Exception as bm25_err:
            logger.info("BM25 搜尋失敗: %s", bm25_err, exc_info=True)

        # ── RRF 合併（k=60）─────────────────────────────────
        # score = α × 1/(k+rank_vector) + (1-α) × 1/(k+rank_bm25)
        RRF_K = 60
        ALPHA = 0.7  # 向量權重

        scores: dict[int, float] = {}
        chunk_map: dict[int, KmChunk] = {}

        for rank, chunk in enumerate(vector_rows, start=1):
            scores[chunk.id] = scores.get(chunk.id, 0.0) + ALPHA / (RRF_K + rank)
            chunk_map[chunk.id] = chunk

        for rank, chunk in enumerate(bm25_rows, start=1):
            scores[chunk.id] = scores.get(chunk.id, 0.0) + (1 - ALPHA) / (RRF_K + rank)
            chunk_map[chunk.id] = chunk

        sorted_ids = sorted(scores, key=lambda cid: scores[cid], reverse=True)
        results = [chunk_map[cid] for cid in sorted_ids[:top_k]]

        logger.info(
            "km_retrieve hybrid: vector=%d bm25=%d merged=%d top_k=%d",
            len(vector_rows), len(bm25_rows), len(scores), top_k,
        )
        return results
    except Exception as e:
        logger.exception("KM 向量搜尋失敗: %s", e)
        return []


def format_km_context(chunks: list[KmChunk], show_source: bool = True) -> str:
    """將 retrieved chunks 格式化為 system prompt 的參考資料字串。"""
    if not chunks:
        return ""

    parts: list[str] = []
    for chunk in chunks:
        if show_source:
            doc_name = chunk.document.filename if chunk.document else "未知文件"
            parts.append(f"--- 來源：{doc_name} ---\n{chunk.content.strip()}")
        else:
            parts.append(chunk.content.strip())

    return "\n\n".join(parts)


# ──────────────────────────────────────────────────────────────────────────────
# FAQ 精確比對模式
# ──────────────────────────────────────────────────────────────────────────────

FAQ_TOP_K = 2                    # 最多回傳幾筆 FAQ 結果


def extract_faq_question(content: str) -> str:
    """從 FAQ chunk 中取出 Q（問題）部分。
    支援格式：Q: / Q： / Question: / 問: / 問題:
    """
    match = _re.search(
        r'^[ \t]*(?:Q[：:]|Question\s*[：:]|問[：:]|問題[：:])\s*(.+)',
        content,
        _re.MULTILINE,
    )
    return match.group(1).strip() if match else ""


def extract_faq_answer(content: str) -> str:
    """從 FAQ chunk 中取出 A（答案）部分的原文。
    支援格式：A: / A： / Answer: / 答: / 答案:
    """
    match = _re.search(
        r'(?:^|\n)[ \t]*(?:A[：:]|Answer\s*[：:]|答[：:]|答案[：:])\s*([\s\S]+)',
        content,
    )
    return match.group(1).strip() if match else content.strip()


def km_faq_retrieve_sync(
    query: str,
    db: Session,
    tenant_id: str,
    user_id: int,
    knowledge_base_id: int,
    top_k: int = 3,
) -> "list[tuple[KmChunk, float]]":
    """FAQ 精確比對：RRF（vector + BM25）。

    索引策略：
      - embedding：Q+A 全文（語意召回）
      - content_tsv（BM25）：Q+A 全文（關鍵字覆蓋，含 tag 與答案術語）

    兩者皆索引全文，RRF 合併分數；語意精準度由 vector（ALPHA=0.7）主導。
    回傳 [(chunk, cosine_similarity), ...] 依 RRF 分數降序，空 list 表示失敗或無結果。
    """
    import sqlalchemy as _sa

    FETCH_K = 20
    RRF_K = 60
    ALPHA = 0.7  # 向量權重（Q-only embedding 已夠精準，維持一般設定）

    # 驗證 KB 存在且屬於當前 tenant，防止跨租戶存取
    kb_exists = db.query(KmKnowledgeBase.id).filter(
        KmKnowledgeBase.id == knowledge_base_id,
        KmKnowledgeBase.tenant_id == tenant_id,
    ).first()
    if not kb_exists:
        logger.warning("FAQ 檢索拒絕：knowledge_base_id=%s 不屬於 tenant_id=%s", knowledge_base_id, tenant_id)
        return []

    embed_params = _get_embed_params(db, tenant_id)
    if not embed_params:
        logger.warning("FAQ 檢索失敗：tenant_id=%s 無 Embedding provider", tenant_id)
        return []

    embed_provider, embed_model, embed_key, embed_base = embed_params
    try:
        vectors, _ = embed_texts_sync(
            [query],
            model=embed_model,
            api_key=embed_key,
            provider=embed_provider,
            api_base=embed_base,
            task_type="retrieval_query",
        )
        if not vectors:
            return []
        query_embedding = vectors[0]
    except Exception as e:
        logger.warning("FAQ embedding 失敗: %s", e)
        return []

    base_filter = [
        KmDocument.knowledge_base_id == knowledge_base_id,
        KmDocument.status == "ready",
    ]

    # ── 1. Vector search ──
    sim_map: dict[int, float] = {}
    chunk_map: dict[int, KmChunk] = {}
    vector_rows = []
    try:
        distance_col = KmChunk.embedding.cosine_distance(query_embedding).label("_dist")
        vector_rows = (
            db.query(KmChunk, distance_col)
            .join(KmDocument, KmChunk.document_id == KmDocument.id)
            .filter(*base_filter)
            .order_by(_sa.text("_dist"))
            .limit(FETCH_K)
            .all()
        )
        for chunk, dist in vector_rows:
            sim_map[chunk.id] = 1.0 - float(dist)
            chunk_map[chunk.id] = chunk
    except Exception as e:
        logger.warning("FAQ 向量搜尋失敗: %s", e)
        return []

    # ── 2. BM25 search ──
    bm25_ids: list[int] = []
    try:
        tsv_raw = db.execute(
            _sa.text("SELECT to_tsvector('public.cjk', :q)"), {"q": query}
        ).scalar()
        if tsv_raw:
            lexemes = [
                f"'{part.split(':')[0].strip(chr(39))}'"
                for part in str(tsv_raw).split()
                if part
            ]
            if lexemes:
                or_tsquery = " | ".join(lexemes)
                tsq = _sa.func.to_tsquery("public.cjk", or_tsquery)
                bm25_rows = (
                    db.query(KmChunk)
                    .join(KmDocument, KmChunk.document_id == KmDocument.id)
                    .filter(*base_filter, KmChunk.content_tsv.op("@@")(tsq))
                    .order_by(_sa.func.ts_rank(KmChunk.content_tsv, tsq, 1).desc())
                    .limit(FETCH_K)
                    .all()
                )
                for chunk in bm25_rows:
                    chunk_map[chunk.id] = chunk
                    bm25_ids.append(chunk.id)
    except Exception as bm25_err:
        logger.info("FAQ BM25 搜尋失敗: %s", bm25_err)

    # ── 2b. 補算 BM25 命中但不在 vector top-20 的 cosine similarity ──
    bm25_only_ids = [cid for cid in bm25_ids if cid not in sim_map]
    if bm25_only_ids:
        try:
            dist_col = KmChunk.embedding.cosine_distance(query_embedding).label("_dist")
            for chunk, dist in (
                db.query(KmChunk, dist_col)
                .filter(KmChunk.id.in_(bm25_only_ids))
                .all()
            ):
                sim_map[chunk.id] = 1.0 - float(dist)
                chunk_map[chunk.id] = chunk
        except Exception as e:
            logger.info("FAQ 補算 BM25-only similarity 失敗: %s", e)

    # ── 3. RRF 合併排序 ──
    scores: dict[int, float] = {}
    for rank, (chunk, _dist) in enumerate(vector_rows, start=1):
        scores[chunk.id] = scores.get(chunk.id, 0.0) + ALPHA / (RRF_K + rank)
    for rank, cid in enumerate(bm25_ids, start=1):
        scores[cid] = scores.get(cid, 0.0) + (1 - ALPHA) / (RRF_K + rank)

    sorted_ids = sorted(scores, key=lambda cid: scores[cid], reverse=True)

    # ── 4. top-k（不設 similarity gate，讓 RRF 分數自然篩選）──
    results: list[tuple[KmChunk, float]] = []
    for cid in sorted_ids:
        sim = sim_map.get(cid, 0.0)
        results.append((chunk_map[cid], sim))
        if len(results) >= top_k:
            break

    logger.info(
        "FAQ retrieve: vector=%d bm25=%d top-%d similarities=%s",
        len(vector_rows), len(bm25_ids), top_k,
        [f"{s:.4f}" for _, s in results],
    )
    return results


async def km_faq_llm_select(
    query: str,
    candidates: "list[tuple[KmChunk, float]]",
    model: str,
    db: Session,
    tenant_id: str,
) -> "tuple[list[tuple[KmChunk, float]], object | None, int]":
    """用 LLM 從候選 FAQ chunks 中選出最相關的。

    候選清單由 km_faq_retrieve_sync 提供（已含 RRF 排序）。
    LLM 只做「選哪幾個」，不生成答案，保證原文 A 零失真。

    回傳 (selected, usage, latency_ms)：
      - 找不到相關 → selected = []
      - LLM 失敗 → fallback：selected = candidates（全部），usage = None，latency_ms = 0
    """
    import time as _time
    import json
    import re

    from app.services.llm_caller import LLMProviderNotConfigured, call_llm

    if not candidates:
        return [], None, 0

    lines: list[str] = []
    for i, (chunk, _) in enumerate(candidates, 1):
        q = extract_faq_question(chunk.content)
        a = extract_faq_answer(chunk.content)
        a_preview = a[:150].replace("\n", " ") + ("…" if len(a) > 150 else "")
        lines.append(f"{i}. Q：{q}\n   A：{a_preview}")

    candidates_text = "\n\n".join(lines)

    from pathlib import Path
    _prompt_file = Path(__file__).resolve().parents[2] / "config" / "system_prompt_faq_direct.md"
    try:
        system_prompt = _prompt_file.read_text(encoding="utf-8").strip()
    except OSError:
        system_prompt = (
            "你是 FAQ 篩選器。根據使用者問題，從候選 FAQ 中選出最相關的項目。\n"
            "只回傳 JSON，格式：{\"selected\": [1, 2]}（編號陣列，1-based）。\n"
            "沒有相關項目時回傳：{\"selected\": []}。\n"
            "不要輸出任何 JSON 以外的文字。"
        )
    user_msg = f"使用者問題：{query}\n\n候選 FAQ：\n{candidates_text}"

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_msg},
    ]

    t0 = _time.perf_counter()
    try:
        answer, usage, latency_ms = await call_llm(
            model=model,
            messages=messages,
            db=db,
            tenant_id=tenant_id,
            temperature=0,
            timeout=30,
        )
        m = re.search(r'\{[^{}]*\}', answer, re.DOTALL)
        if m:
            data = json.loads(m.group())
            selected_indices = [int(x) for x in data.get("selected", [])]
            selected = [
                candidates[i - 1]
                for i in selected_indices
                if 1 <= i <= len(candidates)
            ]
            logger.info(
                "FAQ LLM select: query=%r candidates=%d selected=%s",
                query, len(candidates), selected_indices,
            )
            return selected, usage, latency_ms
        logger.warning("FAQ LLM select: 無法解析 JSON，fallback。原始回應: %r", answer)
        return candidates, usage, latency_ms
    except LLMProviderNotConfigured:
        logger.warning("FAQ LLM select: LLM 未設定，fallback 到全部候選")
    except Exception as e:
        logger.warning("FAQ LLM select 失敗，fallback 到全部候選: %s", e)

    elapsed = int((_time.perf_counter() - t0) * 1000)
    return candidates, None, elapsed  # fallback：LLM 失敗時直接回傳全部候選


def log_km_query(
    db: Session,
    *,
    tenant_id: str,
    user_id: int | None,
    knowledge_base_id: int,
    answer_mode: str,
    query: str,
    hit: bool,
    matched_chunk_ids: list[str],
    session_type: str = "internal",
    widget_session_id: str | None = None,
    chat_thread_id: str | None = None,
) -> None:
    """記錄一次 KB 查詢結果，供零命中統計與知識庫品質分析使用。

    hit=False 代表零命中（RAG 無相關 chunks / direct 無 LLM 選取結果）。
    此函式設計為「靜默失敗」—— 任何例外都只 warning log，不影響主流程。
    """
    import uuid as _uuid
    try:
        row = KmQueryLog(
            tenant_id=tenant_id,
            user_id=user_id,
            knowledge_base_id=knowledge_base_id,
            answer_mode=answer_mode,
            query=query,
            hit=hit,
            matched_chunk_ids=matched_chunk_ids or [],
            session_type=session_type,
            widget_session_id=widget_session_id or None,
            chat_thread_id=_uuid.UUID(str(chat_thread_id)) if chat_thread_id else None,
        )
        db.add(row)
        db.flush()
    except Exception as e:
        logger.warning("log_km_query 寫入失敗（不影響主流程）: %s", e)


def log_bot_query(
    db: Session,
    *,
    tenant_id: str,
    bot_id: int,
    session_id: str | None,
    query: str,
    hit: bool,
) -> None:
    """記錄一次 Bot Widget 查詢結果，供零命中統計與 Bot 品質分析使用。

    此函式設計為「靜默失敗」—— 任何例外都只 warning log，不影響主流程。
    """
    try:
        from app.models.bot_query_log import BotQueryLog
        row = BotQueryLog(
            tenant_id=tenant_id,
            bot_id=bot_id,
            session_id=session_id or None,
            query=query,
            hit=hit,
        )
        db.add(row)
        db.flush()
    except Exception as e:
        logger.warning("log_bot_query 寫入失敗（不影響主流程）: %s", e)
